from docx import Document
from typing import Dict, Any, List, Optional
import os
from datetime import datetime
import re
from app.models.citation import StructuredContent, SourceLocation, SourceType

class WordReader:
    def __init__(self, file_path: str):
        self.file_path = file_path
        self.document = None
        
    def open_document(self) -> bool:
        """Open the Word document"""
        try:
            self.document = Document(self.file_path)
            return True
        except Exception as e:
            print(f"Error opening Word document: {str(e)}")
            return False
    
    def get_full_text(self) -> str:
        """Get full text content of the document"""
        if not self.document:
            return ""
            
        paragraphs = []
        for paragraph in self.document.paragraphs:
            if paragraph.text.strip():
                paragraphs.append(paragraph.text)
                
        return "\n".join(paragraphs)
    
    def get_tables(self) -> List[List[List[str]]]:
        """Get all tables from the document"""
        if not self.document:
            return []
            
        tables = []
        for table in self.document.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
        return tables
    
    def get_headers_footers(self) -> Dict[str, str]:
        """Get headers and footers"""
        if not self.document:
            return {}
            
        headers = []
        footers = []
        
        for section in self.document.sections:
            if section.header.text.strip():
                headers.append(section.header.text)
            if section.footer.text.strip():
                footers.append(section.footer.text)
                
        return {
            'headers': headers,
            'footers': footers
        }
    
    def get_metadata(self) -> Dict[str, Any]:
        """Get document metadata"""
        if not self.document:
            return {}
            
        core_properties = self.document.core_properties
        
        metadata = {
            'author': core_properties.author,
            'created': core_properties.created.isoformat() if core_properties.created else None,
            'modified': core_properties.modified.isoformat() if core_properties.modified else None,
            'last_modified_by': core_properties.last_modified_by,
            'title': core_properties.title,
            'subject': core_properties.subject,
            'keywords': core_properties.keywords,
            'category': core_properties.category,
            'comments': core_properties.comments,
            'file_size': os.path.getsize(self.file_path),
            'paragraph_count': len(self.document.paragraphs),
            'table_count': len(self.document.tables)
        }
        
        return {k: v for k, v in metadata.items() if v is not None}
    
    def get_structured_content(self) -> List[StructuredContent]:
        """Get content with source location tracking (paragraph-based for Word)"""
        structured_content = []
        
        if not self.document:
            return structured_content
        
        # Track sections by looking for heading styles
        current_section = "Document Start"
        section_counter = 0
        
        for para_num, paragraph in enumerate(self.document.paragraphs, 1):
            if not paragraph.text.strip():
                continue
                
            # Check if this paragraph is a heading
            if self._is_heading(paragraph):
                section_counter += 1
                current_section = f"Section {section_counter}: {paragraph.text.strip()}"
                
                # Add the heading as a section source
                source_location = SourceLocation(
                    type=SourceType.SECTION,
                    reference=current_section,
                    text=paragraph.text.strip()
                )
                structured_content.append(StructuredContent(
                    content=paragraph.text.strip(),
                    source_location=source_location,
                    metadata={"paragraph_number": para_num, "is_heading": True}
                ))
            else:
                # Regular paragraph content
                source_location = SourceLocation(
                    type=SourceType.PARAGRAPH,
                    reference=f"paragraph {para_num}",
                    text=paragraph.text.strip()[:200] + "..." if len(paragraph.text.strip()) > 200 else paragraph.text.strip()
                )
                structured_content.append(StructuredContent(
                    content=paragraph.text.strip(),
                    source_location=source_location,
                    metadata={
                        "paragraph_number": para_num, 
                        "current_section": current_section,
                        "is_heading": False
                    }
                ))
        
        # Process tables separately
        for table_num, table in enumerate(self.document.tables, 1):
            table_text = self._extract_table_text(table)
            if table_text:
                source_location = SourceLocation(
                    type=SourceType.SECTION,
                    reference=f"Table {table_num}",
                    text=table_text[:200] + "..." if len(table_text) > 200 else table_text
                )
                structured_content.append(StructuredContent(
                    content=table_text,
                    source_location=source_location,
                    metadata={"table_number": table_num, "is_table": True}
                ))
        
        return structured_content
    
    def _is_heading(self, paragraph) -> bool:
        """Check if a paragraph is a heading based on style"""
        try:
            # Check for heading styles
            if paragraph.style.name.startswith('Heading'):
                return True
            
            # Check for common heading patterns
            text = paragraph.text.strip()
            if not text:
                return False
                
            # Look for patterns like "1.", "A.", "Section 1", etc.
            heading_patterns = [
                r'^\d+\.\s*[A-Z]',  # "1. SOMETHING"
                r'^[A-Z]\.\s*[A-Z]',  # "A. SOMETHING"
                r'^SECTION\s+\d+',  # "SECTION 1"
                r'^ARTICLE\s+\d+',  # "ARTICLE 1"
                r'^SCHEDULE\s+[A-Z0-9]+',  # "SCHEDULE A"
                r'^EXHIBIT\s+[A-Z0-9]+',  # "EXHIBIT A"
                r'^APPENDIX\s+[A-Z0-9]+',  # "APPENDIX A"
            ]
            
            for pattern in heading_patterns:
                if re.match(pattern, text.upper()):
                    return True
                    
            # Check if text is all caps and short (likely a heading)
            if text.isupper() and len(text) < 100 and len(text.split()) <= 10:
                return True
                
            return False
        except:
            return False
    
    def _extract_table_text(self, table) -> str:
        """Extract text from a table"""
        table_text = []
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                table_text.append(" | ".join(row_text))
        return "\n".join(table_text)
    
    def get_content_with_citations(self) -> str:
        """Get full text with embedded citation markers"""
        structured_content = self.get_structured_content()
        content_with_citations = []
        
        for content in structured_content:
            if content.metadata and content.metadata.get("is_heading"):
                citation_marker = f"[{content.source_location.reference}]"
            else:
                citation_marker = f"[{content.source_location.reference}]"
            content_with_citations.append(f"{citation_marker} {content.content}")
        
        return "\n\n".join(content_with_citations) 